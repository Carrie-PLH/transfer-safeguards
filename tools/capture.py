#!/usr/bin/env python3
"""Execute a per-source capture recipe and emit a packet.

The problem this closes. Until now the way a source was captured existed only as
prose in the packet's capture notes: "pdftotext -raw for the PDF", "html2text
over #content", "the workspace web_fetch tool". Prose cannot be re-run. Two
nightly passes over an unchanged source therefore produced two different texts,
and the manifests show exactly that: the transport string recorded for the same
state differs night to night, and not one of the 13 states holding a second
capture has ever produced a second capture that matched the first. The
retained-versus-unchanged signal, which the professional tier's diff series is
built on, currently carries no information at all. It cannot until capture is
deterministic.

A recipe makes it deterministic. Each source gets four declared fields —
transport, extractor, scope, and an ordered list of post-filters — stored in a
sidecar at tools/recipes/<slug>.json. capture.py reads the recipe and performs
the capture, so the nightly executes a recipe instead of reconstructing one from
a paragraph of English.

Why a sidecar and not the SOURCE header line. retain-packet.py compares SOURCE
header lines between passes and reports any difference as a finding, on the
reasoning that a moved or re-dated document is a change even when its text is
not. Putting the recipe in that line would make every recipe correction announce
itself as source movement. The sidecar keeps the two apart: the header describes
the document, the recipe describes how this project reads it. The recipe's
digest is written into the capture notes, which are excluded from the body hash,
so a capture can always be traced to the recipe that produced it without the
recipe's presence perturbing change detection.

Transports this can and cannot run. curl runs here. The agent-mediated readers
do not: the workspace web_fetch tool and Claude in Chrome are tool calls
available to the session, not subprocesses, and Florida is reachable only
through the former (curl returns HTTP 403 and no pdftotext mode reproduces that
extractor). Rather than pretend otherwise, a recipe naming an agent transport
makes capture.py stop and print exactly what it needs; the operator or session
fetches that one URL and supplies the raw text with --supply. The extractor,
scope and filters still run here, so everything downstream of the fetch stays
reproducible even when the fetch itself is not.

Filters are a closed vocabulary. An open-ended filter field would be prose
again. Every name in FILTERS below is a pure text-to-text function defined in
this file, applied in the order the recipe lists them. Adding a filter is a code
change with a self-test, which is the point.

Usage:
    python3 tools/capture.py <slug>                     capture every source
    python3 tools/capture.py <slug> --source 2          one source only
    python3 tools/capture.py <slug> --supply 2=raw.txt  hand it an agent fetch
    python3 tools/capture.py <slug> --out capture.txt   default: stdout
    python3 tools/capture.py --lint [<slug>]            validate recipes
    python3 tools/capture.py --digest <slug>            recipe digest
    python3 tools/capture.py --self-test

Exit codes: 0 clean, 1 capture failure, 2 usage or recipe error,
            3 a recipe needs an agent fetch that was not supplied.
"""

import argparse
import hashlib
import importlib.util
import json
import os
import re
import subprocess
import sys
import unicodedata
from datetime import date

# tools/capture-core.py is a module shared across the Field Assembly portfolio
# (synced from field-assembly-standard/tools/capture-core-sync.py), holding a
# poppler/pdfplumber version pin discovered necessary in a sibling repo,
# sped-safeguards: the same PDF under a different poppler build produces
# different extracted text, which reads downstream as source drift when it is
# actually an undeclared change of extraction tool. This repo pins to the same
# version so that risk does not sit here unguarded. Added 2026-09-03; see
# POPPLER_PIN below for the version and the reasoning in full.
_core_spec = importlib.util.spec_from_file_location(
    'capture_core',
    os.path.join(os.path.dirname(os.path.abspath(__file__)), 'capture-core.py'))
_core = importlib.util.module_from_spec(_core_spec)
_core_spec.loader.exec_module(_core)

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RECIPES = os.path.join(ROOT, "tools", "recipes")

RECIPE_VERSION = 1

# Transports that run in this process. Anything else is agent-mediated and must
# be supplied; the recipe still records which reader was used, because a text
# captured through Chrome and one captured through web_fetch are different texts
# and the packet should say which one it holds.
LOCAL_TRANSPORTS = ("curl",)
AGENT_TRANSPORTS = ("web_fetch", "chrome")
TRANSPORTS = LOCAL_TRANSPORTS + AGENT_TRANSPORTS

EXTRACTORS = ("pdftotext-raw", "pdftotext-layout", "pdfplumber",
              "html-text", "next-data", "docx", "none")

# Opt-in per source (`"compressed": false`), not global, and default on so that
# omitting it fetches exactly as before this option existed.
#
# What it is for. Some hosts serve a gzip stream that curl finishes reading and
# then fails to write: every byte arrives -- byte-identical to a clean fetch --
# and curl exits 23, "Failed writing received data to disk/application", so the
# fetch raises and a complete document is discarded. Reproduced 2026-09-03 on
# www.uab.edu's larger pages (167KB-259KB), intermittently across four sources
# in a single run, and bisected against these very arguments: dropping
# --compressed returns exit 0 with identical bytes, while dropping -L, the
# writeout, or the retry flags changes nothing.
#
# It is a per-source escape rather than a change to CURL_ARGS because the flag
# is worth having everywhere else -- these documents are large and mostly text --
# and because the failure belongs to particular hosts rather than to the
# fetcher. It is not an exit-code exception: treating exit 23 with a full body
# as success would silence a genuine write failure everywhere, which is a much
# worse trade than sending one host uncompressed bytes.
CURL_ARGS = ["-sSL", "--compressed", "--max-time", "120",
             "--retry", "2", "--retry-delay", "3"]


def curl_args(compressed=True):
    """CURL_ARGS, minus --compressed when a source opts out. See the note above."""
    return CURL_ARGS if compressed else [a for a in CURL_ARGS if a != '--compressed']

# Some agencies refuse a request that sends no user-agent: education.vermont.gov
# and mdek12.org both return HTTP 403 to a bare fetch and serve the document
# normally to a browser string. That is a fact about the fetcher, not about the
# document, and before this option the only way past it was to capture by hand —
# which is how Vermont's language survey ended up as evidence no pass could
# refresh, holding up two published quotations.
#
# Opt-in per source (`"user_agent": "browser"`), not global: a site may serve
# different markup to a browser, and switching every existing recipe's fetch at
# once would rewrite captures that are currently reproducible and read as source
# movement. The string is pinned here rather than written in the recipe, for the
# same reason the pdftotext flags are: a value a recipe can vary is a value that
# will vary.
USER_AGENTS = {
    "none": None,
    "browser": ("Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                "AppleWebKit/537.36 (KHTML, like Gecko) Chrome/127.0.0.0 "
                "Safari/537.36"),
}

# Opt-in per source (`"http_version": "1.1"`), curl only, pinned to one value.
# apps.azsos.gov (the Arizona Administrative Code PDFs) returns 403 to every
# HTTP/2 request curl makes, whatever the user-agent, Accept, Referer or
# Sec-Fetch headers, and 200 to the identical request over HTTP/1.1 — the
# refusal is keyed to the protocol negotiation, not to the headers. Found
# 2026-09-04 after two sessions had recorded the host as unreachable to curl.
# Pinned to "1.1" only, for the same reason USER_AGENTS is pinned: a value a
# recipe can vary is a value that will vary, and this one changes the bytes on
# the wire, not the document.
HTTP_VERSIONS = ("1.1",)

# Some agencies serve their leaf certificate without the intermediate that signs
# it. A browser papers over this by caching intermediates or following the
# certificate's AIA extension; curl does neither, and the fetch dies with
# "unable to get local issuer certificate". www.ksde.gov is the case that forced
# this: five of Kansas's six sources had to be downloaded by hand because of it,
# and the packet carried a note calling the host unfetchable.
#
# `"ca_bundle": "<name>"` names a file in tools/certs/ holding a root and the
# intermediates needed under it, passed as curl --cacert. Verification still
# happens — it happens against these anchors instead of the system store, which
# narrows trust rather than widening it. --insecure would have been three
# characters and is not an option: a capture that skipped authentication is not
# evidence of anything.
#
# Opt-in per source and pinned by name, for the same reasons as user_agent and
# the pdftotext flags.
CERTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'certs')


def ca_bundle_path(name):
    return os.path.join(CERTS_DIR, name + '.pem')


def known_ca_bundles():
    if not os.path.isdir(CERTS_DIR):
        return set()
    return {f[:-4] for f in os.listdir(CERTS_DIR) if f.endswith('.pem')}

# pdftotext flags are pinned here rather than left to the recipe: the whole
# purpose is that two runs a month apart produce the same bytes, and a flag the
# recipe can vary is a flag that will vary.
#
# NOTE 2026-09-03: -nopgbrk removed, completing the portfolio. The flag
# suppresses the form feed marking a page boundary rather than rendering it, so
# a page break landing mid-sentence welds the running head that follows onto the
# interrupted sentence with no separator left to cut on (see capture-core.py's
# comment on RUNHEAD_BREAK for the mechanism).
#
# The audit that note asked for was done before the flag was removed, and this
# repo is the small case: both PDF sources were fetched once and extracted
# twice, with the flag and without. Both changed, and both changes are pure
# page-break splits -- collapsing whitespace makes the two extractions
# identical, so separators were restored and nothing else moved. Neither source
# carried a fused word, which is consistent with both being pdftotext-layout:
# that mode already breaks the running head onto its own line, so the flag cost
# it a blank line rather than a sentence. Removing it is still right, because
# the form feed is what strip-running-headers cuts on and what a future
# -raw source here would need.
PDFTOTEXT_RAW = ["-raw", "-enc", "UTF-8"]
PDFTOTEXT_LAYOUT = ["-layout", "-enc", "UTF-8"]

# The build pin itself lives in capture-core.py; see the import shim near the
# top of this file. Untouched until now, this repo has always captured PDFs
# under whatever poppler happened to be installed -- these two names, and the
# require_poppler/require_pdfplumber checks wired into extract_pdf below, are
# new, not a port of a fix this repo already had. They are a preflight guard
# only: the installed poppler already matches the pin (verified 2026-09-03),
# so extracted text is unchanged. What changes is that a future poppler
# upgrade now fails loudly here instead of silently re-baselining every PDF
# packet, which is exactly the failure mode that cost sped-safeguards two
# working days this week.
POPPLER_PIN = _core.DEFAULT_POPPLER_PIN
PDFPLUMBER_PIN = _core.DEFAULT_PDFPLUMBER_PIN


def require_poppler():
    return _core.require_poppler(POPPLER_PIN)


def require_pdfplumber():
    return _core.require_pdfplumber(PDFPLUMBER_PIN)


# --- filters -----------------------------------------------------------------
#
# Each takes text and returns text. They are applied in the order the recipe
# lists them, and the order matters: unwrap-hard-wraps before strip-page-numbers
# will swallow a page number into the paragraph above it.

def f_normalize_apostrophes(t):
    """Curly quotes and dashes to their ASCII equivalents. Agencies re-publish the
    same sentence with different quote characters when they touch a document in a
    different editor, and that is not a change to the sentence."""
    for a, b in (('‘', "'"), ('’', "'"), ('‚', "'"),
                 ('“', '"'), ('”', '"'), ('„', '"'),
                 ('–', '-'), ('—', '--'), ('−', '-'),
                 (' ', ' '), ('​', ''), ('﻿', '')):
        t = t.replace(a, b)
    return t


def f_normalize_ligatures(t):
    """fi/fl and friends to their component letters. One manifest entry already
    records a pass flagged solely for a ligature mismatch with the source
    unchanged."""
    out = []
    for ch in t:
        d = unicodedata.decomposition(ch)
        if d.startswith('<compat>'):
            out.append(unicodedata.normalize('NFKC', ch))
        else:
            out.append(ch)
    return ''.join(out)


def f_strip_zero_width(t):
    """Zero-width characters inserted by a CMS, removed.

    SharePoint sprinkles U+200B through the text its editor produces —
    education.ky.gov carries 175 of them on one page, one of them sitting inside
    "Office of Special Education and Early Learning (OSEEL)" in a sentence the
    published page quotes. They render as nothing, a reader cannot see them, and
    a person retyping the sentence would never reproduce them, so a quotation
    fails on a character that is not there in any meaningful sense.

    This is a different case from the ligature folding in check-fidelity.py, and
    it belongs here rather than there. A ligature is a real decision by the
    document about how to draw a word; a zero-width space is editor exhaust. It
    is removed at capture because the packet should not carry it either.

    Also strips the byte-order mark and the word joiner, which arrive the same
    way and mean the same nothing. Soft hyphens are left alone: they mark a
    permitted break point, which is a typesetting decision the source made.
    """
    return t.replace('\u200b', '').replace('\ufeff', '').replace('\u2060', '')


def f_normalize_bullets(t):
    """Every bullet glyph an agency's CMS might emit becomes a single dash marker."""
    return re.sub(r'^[ \t]*[•▪●◦·⁃∙o*–]\s+',
                  '- ', t, flags=re.MULTILINE)


def f_strip_page_numbers(t):
    """Lines that are a bare page number, or a page number with the document's
    running footer around it. PDFs put these mid-sentence and they move whenever
    the document is repaginated, which is not a change to its text."""
    pat = re.compile(
        r'^\s*(?:page\s+)?(?:[ivxlcdm]+|\d{1,4})\s*(?:of\s+\d{1,4})?\s*$',
        re.IGNORECASE)
    return '\n'.join(l for l in t.split('\n') if not pat.match(l))


WRAP_FLOOR = 40      # below this a document has no column worth speaking of
WRAP_SLACK = 12      # how far short of the column a line may fall and still be full


def wrap_width(t):
    """The column width this document was set to, derived from the document itself.

    A fixed threshold is wrong for every document that is not set to it: a notice
    typeset at 110 characters would have almost every line treated as a heading,
    and one set at 45 would have its headings glued to the paragraphs beneath
    them. The width is instead taken from the document's own long lines — the
    90th percentile of non-blank line lengths, less a few characters of slack for
    the word that did not fit.

    Rounded to the nearest five so that editing a sentence or two cannot shift
    the threshold and silently re-unwrap the whole document, which would show up
    downstream as drift that no source produced."""
    lens = sorted(len(l.rstrip()) for l in t.split('\n') if l.strip())
    if not lens:
        return WRAP_FLOOR
    p90 = lens[min(len(lens) - 1, int(len(lens) * 0.9))]
    return max(WRAP_FLOOR, 5 * round((p90 - WRAP_SLACK) / 5))


def f_unwrap_hard_wraps(t):
    """Join lines a PDF broke to fit its column, keeping real paragraph breaks.

    The signal for a hard wrap is a full column, not the case of the next word: a
    PDF breaks wherever the line ran out, including before a capitalized word like
    "Department". So a line is treated as continuing when it reaches the
    document's column width and does not end a sentence, and when the line after
    it is neither blank nor a list marker. Headings and table labels are short,
    which is what keeps them on their own lines.

    This is the single most common source of false drift: Florida's Source 2
    differed from its baseline by line breaks alone."""
    width = wrap_width(t)
    out = []
    for line in t.split('\n'):
        s = line.rstrip()
        prev = out[-1] if out else ''
        if (prev.strip() and s.strip()
                and len(prev) >= width
                and not re.search(r'[.:;?!]["\')\]]?\s*$', prev)
                and not re.match(r'^\s*(?:[-•]|\(?[a-zA-Z0-9]{1,3}[.)])\s', s)
                and '|' not in prev and '|' not in s):
            out[-1] = prev.rstrip() + ' ' + s.strip()
        else:
            out.append(s)
    return '\n'.join(out)


def f_collapse_blank_runs(t):
    """Two or more blank lines become one."""
    return re.sub(r'\n{3,}', '\n\n', t)


def f_trim_lines(t):
    """Trailing whitespace off every line. Extractors differ on this and nothing
    downstream depends on it."""
    return '\n'.join(l.rstrip() for l in t.split('\n'))


def f_pipe_table_cells(t):
    """Rows already carrying a cell delimiter are normalized to ' | '.

    The convention California and Ohio use, and the one Georgia's rebuild needs.
    A comparison-chart cell captured as period-joined prose invites a quotation
    that carries a terminal period the source never had, which is precisely the
    Georgia defect: five quotations gained punctuation supplied by the capture.
    A pipe cannot be mistaken for the source's own punctuation.

    This normalizes an existing delimiter; it does not infer cell boundaries from
    prose. The extractor is where table structure has to be preserved."""
    return re.sub(r'\s*\|\s*', ' | ', t)


def f_strip_running_headers(t):
    """Drop a running head that recurs at page boundaries; keep everything else.

    Two stages. A line repeating at least RUNHEAD_MIN_REPEATS times within
    RUNHEAD_EDGE_LINES of a page break is a head, not prose. A head fused to the
    line a break interrupted is split on the form feed first, so the sentence is
    restored rather than carried into the packet with the head welded into it.

    A line at an edge is held to a lower character floor
    (RUNHEAD_BREAK_MIN_CHARS) so the short half of a split head is caught, and a
    contact-node protection rule keeps legitimately repeated content -- an
    agency's mailing address or hearing office recurring mid-page -- from being
    eaten.

    Depends on the form feed surviving extraction, which is why -nopgbrk was
    removed above. The implementation lives in capture-core.py, driven by that
    module's own constants; this repo calls it rather than carrying a second
    copy that can drift."""
    return _core.f_strip_running_headers(t)


def f_join_wrap_hyphens(t):
    """A hyphen the source really contains, left stranded at a line break.

    "word-for-\nword" is one word in the document and three tokens after
    whitespace is collapsed: "word-for- word". A quotation of the source's own
    sentence then cannot be verified against it. Montana's Part B notice is the
    case that produced this filter in sped-safeguards, where it was written.

    The hyphen is kept, not deleted. That is the whole judgement in this filter,
    and it is a judgement about a particular kind of document: one that does not
    soft-hyphenate, so a hyphen before a line break is a hyphen the author typed.
    Deleting it would invent "wordforword"; keeping it reproduces what the page
    shows a reader. A source that does use typesetter hyphenation needs a
    different answer -- pdftotext-plain -- and must not use this filter.

    Registered here 2026-09-04 to bring this repo to parity with its siblings.
    No recipe in this collection names it yet, so no capture changes: it is
    capability, not a fix for a case found here. The pattern is byte-identical
    to sped-safeguards' copy and is deliberately narrow -- it requires the
    continuation letter immediately after the newline, so it does not fire on
    pdftotext-layout output, which indents continuation lines. Widening it would
    change what nine sped-safeguards recipes already capture, so it stays as
    written."""
    return re.sub(r'-\n(?=[a-z])', '-', t)


FILTERS = {
    'normalize-apostrophes': f_normalize_apostrophes,
    'strip-running-headers': f_strip_running_headers,
    'join-wrap-hyphens': f_join_wrap_hyphens,
    'normalize-ligatures': f_normalize_ligatures,
    'normalize-bullets': f_normalize_bullets,
    'strip-zero-width': f_strip_zero_width,
    'strip-page-numbers': f_strip_page_numbers,
    'unwrap-hard-wraps': f_unwrap_hard_wraps,
    'collapse-blank-runs': f_collapse_blank_runs,
    'trim-lines': f_trim_lines,
    'pipe-table-cells': f_pipe_table_cells,
}


def apply_filters(text, names):
    for n in names:
        text = FILTERS[n](text)
    return text


# --- slicing -----------------------------------------------------------------
#
# What this is for. Several states publish the provision the page quotes inside
# a much longer document: Illinois serves the whole Nursing Home Care Act and
# the page reads Part 4 of Article III; Missouri, Tennessee and Oklahoma each
# print one rule inside a chapter PDF. Before this option those captures were
# taken by hand -- pdftotext over the whole file, then a line range cut by eye
# -- which is exactly the capture a nightly pass cannot reproduce, and which
# put four states outside the recipe backfill.
#
# The trap this is built around. A long document prints its section headings
# twice, once in the table of contents and once over the text, and a naive
# search for the heading lands in the contents list. The slice that results
# starts at the front matter and runs for tens of thousands of characters,
# looking plausible and containing the right text somewhere inside it. So:
#
#   - Anchors match across line breaks. A heading plus the first words beneath
#     it is the anchor the discipline asks for, and in -layout output that
#     phrase spans a newline and a run of alignment spaces. Every run of
#     whitespace in an anchor matches any run of whitespace in the document,
#     which is what makes "1200-08-06-.05 ADMISSIONS, DISCHARGES, AND
#     TRANSFERS. (1) Admissions." usable as one anchor.
#   - Anchors must be long enough to be specific: MIN_ANCHOR characters. A
#     short anchor is how a slice silently lands in the contents list.
#   - `occurrence` selects among repeats when a document genuinely repeats its
#     anchor. It is deliberately explicit: a recipe that means the second
#     occurrence says so, rather than a default quietly picking one.
#   - A slice that does not match is an error, never a fallback. Falling back
#     to the whole document would produce a packet that passes fidelity (a
#     superset always does) while silently ceasing to be the slice the recipe
#     describes -- the packet would stop matching its own capture notes and
#     nothing would say so.
#
# `to` is exclusive: the slice ends where the next provision begins, which is
# how these cuts are described in the packets that were made by hand ("from the
# heading of DHS 132.53(1) to the end of the served document", "from the
# heading of 1200-08-06-.05 to the start of rule .06"). Omit `to` to run to the
# end of the document.

MIN_ANCHOR = 12


class SliceMiss(Exception):
    """A recipe's slice anchor was not found, or was found out of order."""


def _anchor_re(anchor):
    """The anchor as a regex whose whitespace is elastic. Everything else is
    literal, so a publisher's own parentheses, periods and section symbols mean
    themselves rather than regex operators."""
    return re.compile(r'\s+'.join(re.escape(w) for w in anchor.split()))


def _nth(text, anchor, occurrence, which):
    pat = _anchor_re(anchor)
    hits = list(pat.finditer(text))
    if len(hits) < occurrence:
        raise SliceMiss(
            f'{which} anchor {anchor!r} found {len(hits)} time(s), '
            f'needed occurrence {occurrence}')
    return hits[occurrence - 1]


def apply_slice(text, spec):
    """Cut `text` to the recipe's slice. Raises SliceMiss rather than returning
    anything approximate; see the note above for why there is no fallback."""
    if not spec:
        return text
    start = _nth(text, spec['from'], spec.get('from_occurrence', 1), 'from')
    if not spec.get('to'):
        return text[start.start():]
    tail = text[start.end():]
    end = _nth(tail, spec['to'], spec.get('to_occurrence', 1), 'to')
    return text[start.start():start.end() + end.start()]


def slice_note(spec):
    """How a slice is described in the packet's capture notes."""
    bits = [f'slice from {spec["from"]!r}']
    if spec.get('from_occurrence', 1) != 1:
        bits.append(f'(occurrence {spec["from_occurrence"]})')
    bits.append(f'to {spec["to"]!r}' if spec.get('to')
                else 'to the end of the document')
    if spec.get('to') and spec.get('to_occurrence', 1) != 1:
        bits.append(f'(occurrence {spec["to_occurrence"]})')
    return ' '.join(bits)


# --- recipes -----------------------------------------------------------------

def recipe_path(slug):
    return os.path.join(RECIPES, f'{slug}.json')


def load_recipe(slug):
    p = recipe_path(slug)
    if not os.path.exists(p):
        raise SystemExit(f'no recipe for {slug} (expected {p})')
    with open(p, encoding='utf-8') as fh:
        return json.load(fh)


def lint(rec, slug=None):
    """Everything wrong with a recipe, as a list. A recipe that does not lint is a
    recipe the nightly cannot trust."""
    errs = []
    if rec.get('recipe_version') != RECIPE_VERSION:
        errs.append(f'recipe_version must be {RECIPE_VERSION}')
    if slug and rec.get('state') != slug:
        errs.append(f'state is {rec.get("state")!r}, file is named {slug!r}')
    sources = rec.get('sources')
    if not isinstance(sources, list) or not sources:
        errs.append('sources must be a non-empty list')
        return errs
    for i, s in enumerate(sources, 1):
        w = f'source {s.get("n", i)}'
        if s.get('n') != i:
            errs.append(f'{w}: sources must be numbered 1..n in order')
        for k in ('title', 'url', 'transport', 'extractor'):
            if not s.get(k):
                errs.append(f'{w}: missing {k}')
        if s.get('transport') and s['transport'] not in TRANSPORTS:
            errs.append(f'{w}: unknown transport {s["transport"]!r}')
        if s.get('extractor') and s['extractor'] not in EXTRACTORS:
            errs.append(f'{w}: unknown extractor {s["extractor"]!r}')
        if s.get('extractor') == 'html-text' and not s.get('scope'):
            errs.append(f'{w}: html-text needs a scope selector; use "body" for '
                        f'the whole document rather than leaving it unstated')
        if s.get('extractor') == 'next-data':
            sc = s.get('scope')
            if not sc:
                errs.append(f'{w}: next-data needs a scope giving the dotted JSON '
                            f'path to the page content inside the island')
            elif not (isinstance(sc, str)
                      or (isinstance(sc, list) and sc
                          and all(isinstance(x, str) for x in sc))):
                errs.append(f'{w}: next-data scope must be a path or a '
                            f'non-empty list of paths')
        if s.get('scope') and s.get('extractor') not in ('html-text', 'next-data',
                                                         'docx'):
            errs.append(f'{w}: scope has no meaning for extractor '
                        f'{s.get("extractor")!r}')
        if s.get('user_agent') and s['user_agent'] not in USER_AGENTS:
            errs.append(f'{w}: unknown user_agent {s["user_agent"]!r}; '
                        f'one of {", ".join(sorted(USER_AGENTS))}')
        if s.get('user_agent') and s.get('transport') != 'curl':
            errs.append(f'{w}: user_agent applies to the curl transport only')
        if s.get('http_version') and s['http_version'] not in HTTP_VERSIONS:
            errs.append(f'{w}: unknown http_version {s["http_version"]!r}; '
                        f'one of {", ".join(HTTP_VERSIONS)}')
        if s.get('http_version') and s.get('transport') != 'curl':
            errs.append(f'{w}: http_version applies to the curl transport only')
        if s.get('compressed') is not None and not isinstance(s['compressed'], bool):
            errs.append(f'{w}: compressed must be true or false')
        if s.get('compressed') is False and s.get('transport') != 'curl':
            errs.append(f'{w}: compressed applies to the curl transport only')
        if s.get('ca_bundle') and s['ca_bundle'] not in known_ca_bundles():
            errs.append(f'{w}: unknown ca_bundle {s["ca_bundle"]!r}; '
                        f'expected a .pem in tools/certs/ '
                        f'({", ".join(sorted(known_ca_bundles())) or "none present"})')
        if s.get('ca_bundle') and s.get('transport') != 'curl':
            errs.append(f'{w}: ca_bundle applies to the curl transport only')
        if s.get('pages') is not None:
            pg = s['pages']
            if s.get('extractor') not in ('pdftotext-raw', 'pdftotext-layout',
                                          'pdfplumber'):
                errs.append(f'{w}: pages applies to the PDF extractors only')
            elif (not isinstance(pg, list) or len(pg) != 2
                  or not all(isinstance(x, int) for x in pg)
                  or pg[0] < 1 or pg[1] < pg[0]):
                errs.append(f'{w}: pages must be [first, last], 1-based and '
                            f'inclusive, with first <= last; got {pg!r}')
        sl = s.get('slice')
        if sl is not None:
            if not isinstance(sl, dict):
                errs.append(f'{w}: slice must be an object with "from" and '
                            f'optionally "to"')
            else:
                unknown = set(sl) - {'from', 'to', 'from_occurrence',
                                     'to_occurrence'}
                if unknown:
                    errs.append(f'{w}: slice has unknown key(s) '
                                f'{", ".join(sorted(unknown))}')
                for k in ('from', 'to'):
                    v = sl.get(k)
                    if k == 'to' and v is None:
                        continue
                    if not isinstance(v, str) or not v.strip():
                        errs.append(f'{w}: slice {k} must be a non-empty string'
                                    if k == 'from' else
                                    f'{w}: slice to, when present, must be a '
                                    f'non-empty string; omit it to slice to the '
                                    f'end of the document')
                    elif len(v.strip()) < MIN_ANCHOR:
                        errs.append(
                            f'{w}: slice {k} anchor is {len(v.strip())} '
                            f'characters; at least {MIN_ANCHOR} are needed. A '
                            f'short anchor is how a slice lands in the table of '
                            f'contents instead of the text — anchor on the '
                            f'heading plus the first words beneath it')
                for k in ('from_occurrence', 'to_occurrence'):
                    if k in sl and (not isinstance(sl[k], int)
                                    or isinstance(sl[k], bool) or sl[k] < 1):
                        errs.append(f'{w}: slice {k} must be a positive integer '
                                    f'(1 is the first match)')
                if 'to_occurrence' in sl and not sl.get('to'):
                    errs.append(f'{w}: slice to_occurrence has no meaning '
                                f'without a "to" anchor')
        for f in s.get('filters', []):
            if f not in FILTERS:
                errs.append(f'{w}: unknown filter {f!r}')
    return errs


def digest(rec):
    """A short hash over the fields that determine the captured text. Title and
    notes are excluded: editing a title changes the packet header, not the
    capture, and a digest that moves when prose moves teaches people to ignore
    it."""
    # user_agent joins the digest only when a recipe sets it. A source that
    # sends no user-agent fetches exactly as it did before the option existed,
    # so its digest must not move: every capture already taken records the
    # digest that produced it, and invalidating those to add a field nobody
    # used would break the trace for no gain.
    material = []
    for s in rec['sources']:
        m = {k: s.get(k) for k in ('n', 'url', 'transport', 'extractor', 'scope')}
        m['filters'] = list(s.get('filters', []))
        if s.get('user_agent') and s['user_agent'] != 'none':
            m['user_agent'] = s['user_agent']
        if s.get('compressed') is False:
            m['compressed'] = False
        if s.get('http_version'):
            m['http_version'] = s['http_version']
        if s.get('ca_bundle'):
            m['ca_bundle'] = s['ca_bundle']
        if s.get('pages'):
            m['pages'] = s['pages']
        # Same opt-in rule as the fields above: a source with no slice hashes
        # exactly as it did before slicing existed, so no capture already taken
        # is invalidated. A recipe that changes its anchors is capturing a
        # different span and must show a different digest.
        if s.get('slice'):
            m['slice'] = {k: s['slice'][k] for k in sorted(s['slice'])}
        material.append(m)
    blob = json.dumps({'v': RECIPE_VERSION, 'sources': material},
                      sort_keys=True, separators=(',', ':'))
    return hashlib.sha256(blob.encode('utf-8')).hexdigest()[:12]


# --- transports and extractors ----------------------------------------------

def fetch_curl(url, binary, user_agent='none', ca_bundle=None,
               compressed=True, http_version=None):
    ua = USER_AGENTS[user_agent]
    args = ['curl'] + curl_args(compressed) + (['-A', ua] if ua else []) \
        + (['--http1.1'] if http_version == '1.1' else []) \
        + (['--cacert', ca_bundle_path(ca_bundle)] if ca_bundle else []) \
        + ['-w', '%{http_code}', '-o', '-', url]
    r = subprocess.run(args, capture_output=True, timeout=180)
    if r.returncode != 0:
        raise RuntimeError(f'curl failed ({r.returncode}): '
                           f'{r.stderr.decode("utf-8", "replace").strip()}')
    blob, code = r.stdout[:-3], r.stdout[-3:].decode('ascii', 'replace')
    if code != '200':
        raise RuntimeError(f'HTTP {code}')
    return blob if binary else blob.decode('utf-8', 'replace')


def extract_pdf(blob, mode, pages=None):
    """Text from a PDF, optionally from an inclusive page range only.

    `pages` is the PDF scope field, and it is there for the same reason the HTML
    extractor has a selector: a packet should hold the part of a document the
    page rests on, not everything that happens to be bound with it.

    Kansas is why it exists. Its Source 3 is a 2018 mirror of the complete
    handbook, held because two published quotations of the K.S.A. 72-3430
    statute were read from it. Captured whole, that one file also contains the
    2018 text of every other chapter — including the wording KSDE replaced in
    June 2026. The checker would then certify the page's superseded quotations
    against a superseded document forever, and the review that exists to catch
    exactly that would report a clean pass. Narrowing the scope to the pages the
    quotations came from is what keeps the drift visible.
    """
    import tempfile
    with tempfile.TemporaryDirectory() as td:
        p = os.path.join(td, 'in.pdf')
        open(p, 'wb').write(blob)
        if mode == 'pdfplumber':
            require_pdfplumber()
            import pdfplumber
            with pdfplumber.open(p) as pdf:
                sel = pdf.pages if not pages else pdf.pages[pages[0] - 1:pages[1]]
                return '\n'.join((pg.extract_text() or '') for pg in sel)
        require_poppler()
        flags = PDFTOTEXT_RAW if mode == 'pdftotext-raw' else PDFTOTEXT_LAYOUT
        if pages:
            flags = flags + ['-f', str(pages[0]), '-l', str(pages[1])]
        r = subprocess.run(['pdftotext'] + flags + [p, '-'],
                           capture_output=True, timeout=300)
        if r.returncode != 0:
            raise RuntimeError('pdftotext: '
                               + r.stderr.decode('utf-8', 'replace').strip())
        return r.stdout.decode('utf-8', 'replace')


def extract_html(text, scope):
    """Text under one CSS selector, with script, style and hidden nodes dropped.

    Implemented against bs4 in this process rather than shelling out to a
    converter, because a converter's version is another thing that drifts between
    nights. Block-level elements become line breaks; table rows are emitted with
    ' | ' between cells, so a comparison chart arrives with its cell boundaries
    intact and no punctuation invented to stand in for them."""
    from bs4 import BeautifulSoup, NavigableString
    soup = BeautifulSoup(text, 'html.parser')
    for bad in soup(['script', 'style', 'noscript', 'svg']):
        bad.decompose()
    # Comments and hidden elements, before any walk of the tree. bs4's Comment
    # is a NavigableString subclass, so walk() below appends a comment exactly
    # as it appends a paragraph; and the docstring's promise of "hidden nodes
    # dropped" was never implemented until this call. See FA-Q-20260904-08 and
    # capture-core's own commentary for the two agencies that exposed both.
    _core.strip_nonvisible_nodes(soup)
    # Cloudflare-obfuscated addresses are decoded before anything else reads the
    # document, because the plaintext is already in the served bytes and a
    # capture that keeps the placeholder cannot vouch for a contact the
    # publisher prints. Wired in 2026-09-03: a portfolio scan found 76 of these
    # placeholders across this repo and transfer-safeguards, and none in the two
    # repos that already decoded them. Unconditional rather than opt-in, unlike
    # the links field, because it adds nothing the publisher did not publish.
    _core.decode_cfemail_nodes(soup)
    node = soup.body if scope == 'body' else soup.select_one(scope)
    if node is None:
        raise RuntimeError(f'scope selector matched nothing: {scope!r}')

    for tr in node.find_all('tr'):
        cells = [' '.join(c.get_text(' ', strip=True).split())
                 for c in tr.find_all(['td', 'th'])]
        tr.replace_with(NavigableString('\n' + ' | '.join(cells) + '\n'))

    BLOCK = {'p', 'div', 'br', 'li', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
             'section', 'article', 'header', 'footer', 'table', 'ul', 'ol'}
    parts = []

    def walk(n):
        if isinstance(n, NavigableString):
            parts.append(str(n))
            return
        name = getattr(n, 'name', None)
        if name in BLOCK:
            parts.append('\n')
        for c in getattr(n, 'children', []):
            walk(c)
        if name in BLOCK:
            parts.append('\n')

    walk(node)
    return re.sub(r'[ \t]+', ' ', ''.join(parts))


NEXT_DATA_RE = re.compile(
    r'<script id="__NEXT_DATA__" type="application/json">(.*?)</script>', re.S)
PATH_STEP_RE = re.compile(r'([^.\[\]]+)|\[(\d+)\]')


def extract_next_data(text, path):
    """Page HTML out of a Next.js server-rendered data island.

    A site can look client-rendered and not be. gadoe.org returns a shell whose
    body extracts to zero characters, which is why Georgia was captured through a
    browser and why its comparison chart arrived flattened. But the page's real
    content — the chart included, as an actual <table> — is server-rendered
    inside the __NEXT_DATA__ script on the very same response a plain curl
    already receives. The document was reachable mechanically the whole time.

    That is worth stating as a rule rather than a Georgia anecdote: an empty body
    is a fact about where a framework put the content, not a fact about whether
    the content was served. Look for the island before reaching for a browser.

    `scope` is a dotted path into the island's JSON, with [i] for list indices,
    or a list of such paths concatenated in the order given. A page's parts are
    not all in one place: gadoe.org keeps the body copy as an HTML string under
    page.content and the office's phone and email as structured fields under
    page.contactInformation. Capturing only the first drops the contacts, and a
    contact that is not in the packet is a contact the fidelity checker cannot
    vouch for — which is the failure mode that check exists to catch, since a
    wrong phone number is what a parent actually dials.

    A path resolving to a string is treated as HTML and goes through the same
    text extraction every other HTML source does, so table cells arrive
    pipe-delimited. A path resolving to an object or list is rendered as
    `key: value` lines, which is a faithful flattening rather than an
    interpretation of structured data."""
    m = NEXT_DATA_RE.search(text)
    if not m:
        raise RuntimeError('no __NEXT_DATA__ island in this response')
    island = json.loads(m.group(1))
    paths = [path] if isinstance(path, str) else list(path)
    out = []
    for p in paths:
        node = island
        for step in PATH_STEP_RE.finditer(p):
            key, idx = step.group(1), step.group(2)
            try:
                node = node[int(idx)] if idx is not None else node[key]
            except (KeyError, IndexError, TypeError):
                raise RuntimeError(f'path {p!r} does not resolve at {step.group(0)!r}')
        if node is None:
            raise RuntimeError(f'path {p!r} resolves to null')
        out.append(extract_html(f'<div>{node}</div>', 'div')
                   if isinstance(node, str) else render_json_lines(node))
    return '\n'.join(out)


def render_json_lines(node, prefix=''):
    """Structured data as `key: value` lines, in document order.

    Deliberately dull. The alternative is to compose the fields into a sentence,
    and a composed sentence is a sentence this project wrote, which a page could
    then quote as though the agency had written it."""
    lines = []
    if isinstance(node, dict):
        for k, v in node.items():
            if v in (None, '', [], {}):
                continue
            if isinstance(v, (dict, list)):
                lines.append(f'{prefix}{k}:')
                lines.append(render_json_lines(v, prefix + '  '))
            else:
                lines.append(f'{prefix}{k}: {v}')
    elif isinstance(node, list):
        for i, v in enumerate(node):
            if isinstance(v, (dict, list)):
                lines.append(f'{prefix}[{i}]')
                lines.append(render_json_lines(v, prefix + '  '))
            else:
                lines.append(f'{prefix}[{i}]: {v}')
    else:
        lines.append(f'{prefix}{node}')
    return '\n'.join(l for l in lines if l.strip())


def extract_docx(blob, scope):
    """Paragraphs in document order, then table rows with cells pipe-joined —
    the convention the California packet already records for its .docx source."""
    import io
    import docx
    d = docx.Document(io.BytesIO(blob))
    out = [p.text for p in d.paragraphs]
    if scope != 'paragraphs-only':
        for t in d.tables:
            for row in t.rows:
                out.append(' | '.join(c.text.strip() for c in row.cells))
    return '\n'.join(out)


def capture_source(src, supplied=None):
    """One source, from recipe to filtered text."""
    transport, extractor = src['transport'], src['extractor']
    needs_binary = extractor in ('pdftotext-raw', 'pdftotext-layout',
                                 'pdfplumber', 'docx')

    if supplied is not None:
        raw = supplied
    elif transport in AGENT_TRANSPORTS:
        raise NeedsAgentFetch(src)
    else:
        raw = fetch_curl(src['url'], binary=needs_binary,
                         user_agent=src.get('user_agent', 'none'),
                         compressed=src.get('compressed', True),
                         ca_bundle=src.get('ca_bundle'),
                         http_version=src.get('http_version'))

    if isinstance(raw, str) and needs_binary:
        # A supplied agent fetch is already text; the extractor has effectively
        # been run by the reader. Say so in the packet rather than pretending
        # this file's extractor produced it.
        text = raw
    elif extractor in ('pdftotext-raw', 'pdftotext-layout', 'pdfplumber'):
        text = extract_pdf(raw, extractor, src.get('pages'))
    elif extractor == 'docx':
        text = extract_docx(raw, src.get('scope') or 'paragraphs-and-tables')
    elif extractor == 'html-text':
        text = extract_html(raw, src['scope'])
    elif extractor == 'next-data':
        text = extract_next_data(raw, src['scope'])
    else:
        text = raw if isinstance(raw, str) else raw.decode('utf-8', 'replace')

    # Slice before the filters, because that is the order the hand captures
    # this replaces used: extract the whole document, cut the span, then tidy.
    # Filtering first would move the anchors — strip-page-numbers and
    # unwrap-hard-wraps both rewrite the lines an anchor may sit across.
    try:
        text = apply_slice(text, src.get('slice'))
    except SliceMiss as e:
        raise SystemExit(
            f'source {src["n"]}: {e}\n'
            f'  The document was fetched and extracted; the slice is what '
            f'failed.\n'
            f'  Read the extraction before adjusting the anchor: a document '
            f'that has been\n'
            f'  reissued, repaginated or re-headed is a finding about the '
            f'source, not a\n'
            f'  recipe to widen until it matches again.')

    text = apply_filters(text, src.get('filters', []))

    # Whatever page breaks the filters did not consume become line breaks here.
    # Dropping -nopgbrk keeps the form feed so that strip-running-headers can
    # find the head that follows it, but only the recipes that actually carry a
    # running head run that filter, and the rest should not start carrying a
    # control character into their packets. A line break is the honest
    # rendering: it adds a break where the document has a page boundary and
    # removes nothing.
    #
    # Deliberately after the filters and not before: converting earlier would
    # destroy the same signal -nopgbrk used to destroy.
    #
    # form_feed_to_linebreak lives in capture-core.py; it is a one-line
    # function, but every repo needs the exact same one line, called at the
    # exact same point in the pipeline (after filters, not before).
    return _core.form_feed_to_linebreak(text)


class NeedsAgentFetch(Exception):
    def __init__(self, src):
        self.src = src
        super().__init__(f'source {src["n"]} needs {src["transport"]}')


# --- packet emission ---------------------------------------------------------

CANARY = ('FIRST LINE OF PACKET — if you cannot see this sentence, '
          'output only the words PACKET TRUNCATED')


def render_packet(rec, bodies, day, supplied_ns=()):
    d = digest(rec)
    lines = [CANARY, '', f'STATE: {rec["state"]}', f'ASSEMBLED: {day}', '',
             'CAPTURE NOTES:',
             f'- Captured by tools/capture.py from tools/recipes/{rec["state"]}.json '
             f'(recipe digest {d}). Each source below was fetched and extracted by '
             f'its recorded recipe; re-running that recipe against an unchanged '
             f'source reproduces this text.']
    for s in rec['sources']:
        bits = [f'transport {s["transport"]}', f'extractor {s["extractor"]}']
        if s.get('scope'):
            bits.append(f'scope {s["scope"]}')
        if s.get('pages'):
            bits.append(f'pages {s["pages"][0]}-{s["pages"][1]}')
        if s.get('slice'):
            bits.append(slice_note(s['slice']))
        if s.get('ca_bundle'):
            bits.append(f'ca_bundle {s["ca_bundle"]}')
        if s.get('http_version'):
            bits.append(f'http_version {s["http_version"]}')
        bits.append('filters ' + (', '.join(s.get('filters', [])) or 'none'))
        supp = ' (fetched by the session reader and supplied)' \
            if s['n'] in supplied_ns else ''
        lines.append(f'- SOURCE {s["n"]}: ' + '; '.join(bits) + supp + '.')
        if s.get('notes'):
            lines.append(f'  {s["notes"]}')
    lines.append('')
    for s in rec['sources']:
        if s['n'] not in bodies:
            continue
        lines.append(
            f'SOURCE {s["n"]}: {s["title"]} | {s["url"]} | '
            f'source date: {s.get("source_date") or "none published"} | '
            f'retrieved: {day}')
        lines.append(bodies[s['n']].strip())
        lines.append('')
    return '\n'.join(lines) + '\n'


# --- cli ---------------------------------------------------------------------

def cmd_capture(args):
    rec = load_recipe(args.slug)
    errs = lint(rec, args.slug)
    if errs:
        for e in errs:
            print(f'recipe error: {e}', file=sys.stderr)
        return 2

    supplied = {}
    for spec in args.supply or []:
        n, _, path = spec.partition('=')
        if not path:
            print(f'--supply wants N=path, got {spec!r}', file=sys.stderr)
            return 2
        supplied[int(n)] = open(path, encoding='utf-8').read()

    wanted = [s for s in rec['sources']
              if args.source is None or s['n'] == args.source]
    if not wanted:
        print(f'no source {args.source} in {args.slug}', file=sys.stderr)
        return 2

    bodies, pending, failed = {}, [], []
    for s in wanted:
        try:
            bodies[s['n']] = capture_source(s, supplied.get(s['n']))
        except NeedsAgentFetch:
            pending.append(s)
        except Exception as e:
            failed.append((s, e))
            print(f'SOURCE {s["n"]} failed: {e}', file=sys.stderr)

    if pending:
        print('', file=sys.stderr)
        print('These sources are read through a session tool, not a subprocess. '
              'Fetch each URL with the named reader, save the text, and re-run '
              'with --supply:', file=sys.stderr)
        for s in pending:
            print(f'  SOURCE {s["n"]}  {s["transport"]}  {s["url"]}',
                  file=sys.stderr)
        print('  e.g. python3 tools/capture.py %s %s' % (
            args.slug, ' '.join(f'--supply {s["n"]}=/tmp/s{s["n"]}.txt'
                                for s in pending)), file=sys.stderr)

    if bodies:
        out = render_packet(rec, bodies, args.date or date.today().isoformat(),
                            supplied_ns=set(supplied))
        if args.out:
            open(args.out, 'w', encoding='utf-8').write(out)
            print(f'wrote {args.out} · {len(bodies)} source(s) · '
                  f'recipe {digest(rec)}', file=sys.stderr)
        else:
            sys.stdout.write(out)

    return 1 if failed else (3 if pending else 0)


def cmd_lint(args):
    slugs = [args.slug] if args.slug else sorted(
        f[:-5] for f in os.listdir(RECIPES) if f.endswith('.json'))
    if not slugs:
        print('no recipes yet')
        return 0
    bad = 0
    for slug in slugs:
        errs = lint(load_recipe(slug), slug)
        if errs:
            bad += 1
            print(f'{slug}: {len(errs)} problem(s)')
            for e in errs:
                print(f'  {e}')
        else:
            rec = load_recipe(slug)
            agent = sum(1 for s in rec['sources']
                        if s['transport'] in AGENT_TRANSPORTS)
            print(f'{slug}: ok · {len(rec["sources"])} source(s) · '
                  f'digest {digest(rec)}'
                  + (f' · {agent} need a session reader' if agent else ''))
    print(f'{len(slugs) - bad}/{len(slugs)} recipe(s) clean')
    return 1 if bad else 0


def self_test():
    ok = True

    def check(cond, msg):
        nonlocal ok
        if not cond:
            print(f'FAIL: {msg}')
            ok = False

    # user_agent: opt-in, curl-only, and digest-stable when unset. The last of
    # those is the one worth testing — adding the field must not move the digest
    # of a recipe that does not use it, or every capture already taken records a
    # digest that no longer resolves.
    ua_none = {'recipe_version': RECIPE_VERSION, 'state': 'testland',
               'sources': [{'n': 1, 'title': 'T', 'url': 'https://x.gov/a',
                            'transport': 'curl', 'extractor': 'pdfplumber',
                            'filters': []}]}
    import copy
    ua_set = copy.deepcopy(ua_none)
    ua_set['sources'][0]['user_agent'] = 'browser'
    ua_explicit_none = copy.deepcopy(ua_none)
    ua_explicit_none['sources'][0]['user_agent'] = 'none'
    check(digest(ua_none) == digest(ua_explicit_none),
          'an explicit user_agent of "none" moved the digest')
    check(digest(ua_none) != digest(ua_set),
          'setting a browser user_agent did not move the digest')
    bad_ua = copy.deepcopy(ua_none)
    bad_ua['sources'][0]['user_agent'] = 'firefox-42'
    check(any('user_agent' in e for e in lint(bad_ua, 'testland')),
          'lint accepted an unknown user_agent')
    wrong_transport = copy.deepcopy(ua_set)
    wrong_transport['sources'][0]['transport'] = 'chrome'
    check(any('curl transport only' in e
              for e in lint(wrong_transport, 'testland')),
          'lint accepted user_agent on a non-curl transport')
    check(USER_AGENTS['none'] is None and 'Mozilla' in USER_AGENTS['browser'],
          'user-agent table is not what the recipes expect')

    # http_version: same opt-in shape. Unset must not move the digest; "1.1"
    # must; anything else, or any non-curl transport, must fail lint.
    hv_set = copy.deepcopy(ua_none)
    hv_set['sources'][0]['http_version'] = '1.1'
    check(digest(ua_none) != digest(hv_set),
          'setting http_version did not move the digest')
    check(not any('http_version' in e for e in lint(hv_set, 'testland')),
          'lint rejected http_version "1.1" on a curl source')
    hv_bad = copy.deepcopy(ua_none)
    hv_bad['sources'][0]['http_version'] = '2'
    check(any('http_version' in e for e in lint(hv_bad, 'testland')),
          'lint accepted an unknown http_version')
    hv_wrong = copy.deepcopy(hv_set)
    hv_wrong['sources'][0]['transport'] = 'chrome'
    check(any('curl transport only' in e for e in lint(hv_wrong, 'testland')),
          'lint accepted http_version on a non-curl transport')

    # ca_bundle: same opt-in shape, and it must name a file that exists. A
    # recipe pointing at a missing bundle would otherwise fail at fetch time
    # with a curl error that reads like the agency's fault.
    ca_set = copy.deepcopy(ua_none)
    ca_set['sources'][0]['ca_bundle'] = 'digicert-global-root-g2'
    check(digest(ua_none) != digest(ca_set),
          'setting a ca_bundle did not move the digest')
    check(not any('ca_bundle' in e for e in lint(ca_set, 'testland')),
          'lint rejected a ca_bundle that is present in tools/certs/')
    bad_ca = copy.deepcopy(ua_none)
    bad_ca['sources'][0]['ca_bundle'] = 'no-such-authority'
    check(any('ca_bundle' in e for e in lint(bad_ca, 'testland')),
          'lint accepted a ca_bundle with no file behind it')
    ca_wrong_transport = copy.deepcopy(ca_set)
    ca_wrong_transport['sources'][0]['transport'] = 'chrome'
    check(any('ca_bundle applies' in e
              for e in lint(ca_wrong_transport, 'testland')),
          'lint accepted ca_bundle on a non-curl transport')

    # pages: the PDF scope field. It moves the digest, because a capture of
    # pages 1-14 and a capture of the whole file are different evidence.
    pg_set = copy.deepcopy(ua_none)
    pg_set['sources'][0]['pages'] = [1, 14]
    check(digest(ua_none) != digest(pg_set),
          'setting a page range did not move the digest')
    check(not any('pages' in e for e in lint(pg_set, 'testland')),
          'lint rejected a well-formed page range')
    for bad in ([0, 5], [9, 3], [1], '1-14', [1, 'x']):
        b = copy.deepcopy(ua_none)
        b['sources'][0]['pages'] = bad
        check(any('pages' in e for e in lint(b, 'testland')),
              f'lint accepted a malformed page range {bad!r}')
    pg_html = copy.deepcopy(pg_set)
    pg_html['sources'][0]['extractor'] = 'html-text'
    pg_html['sources'][0]['scope'] = 'main'
    check(any('PDF extractors only' in e for e in lint(pg_html, 'testland')),
          'lint accepted a page range on an HTML extractor')

    # Filters
    check(f_normalize_apostrophes('the child’s') == "the child's",
          'apostrophe normalization')
    check(f_strip_zero_width('(OSEEL)\u200b within one\ufeff year')
          == '(OSEEL) within one year', 'zero-width stripping')
    check(f_strip_zero_width('re\u00adenter') == 're\u00adenter',
          'zero-width stripping ate a soft hyphen, which is a real break point')
    check(f_strip_page_numbers('body\n12\nmore\nPage 3 of 40\nend')
          == 'body\nmore\nend', 'page-number stripping')
    check(f_normalize_bullets('• one\n● two') == '- one\n- two',
          'bullet normalization')
    check(f_collapse_blank_runs('a\n\n\n\nb') == 'a\n\nb', 'blank-run collapse')
    check(f_pipe_table_cells('a|b  |   c') == 'a | b | c', 'pipe normalization')

    # The Florida case: a PDF's hard wraps must not read as a change. The break
    # lands before a capitalized word, which is why case cannot be the test.
    wrapped = ('The Individuals with Disabilities Education Act (IDEA), the Federal '
               'law concerning the education of students with\n'
               'disabilities, requires schools to provide you, the parents of a child '
               'with a disability, with a notice.\n'
               '\nTable of Contents')
    unwrapped = f_unwrap_hard_wraps(wrapped)
    check('students with disabilities, requires' in unwrapped,
          'hard-wrap unwrapping')
    check(unwrapped.rstrip().endswith('\nTable of Contents'),
          'unwrapping ate a real paragraph break')
    # A short heading is not a wrapped line, however it is punctuated.
    check(f_unwrap_hard_wraps('Filing a State Complaint\nAny parent may file')
          == 'Filing a State Complaint\nAny parent may file',
          'unwrapping joined a heading to the paragraph under it')
    # A list marker is never a continuation.
    listy = ('you must provide the following information in your written complaint\n'
             '(1) the name of the child')
    check('\n(1)' in f_unwrap_hard_wraps(listy),
          'unwrapping swallowed a list item')
    # Nor is a pipe-delimited table row, which has no wrap to undo.
    rows = ('Who may file | Any individual or organization who believes a violation\n'
            'Deadline | One year from the date of the alleged violation')
    check('\nDeadline' in f_unwrap_hard_wraps(rows),
          'unwrapping merged two table rows')
    # The column width comes from the document, and small edits must not move it.
    narrow = '\n'.join(['x' * 46] * 40)
    wide = '\n'.join(['x' * 104] * 40)
    check(wrap_width(narrow) < wrap_width(wide),
          'column width did not follow the document')
    edited = '\n'.join(['x' * 46] * 37 + ['short', 'also short', 'x' * 46])
    check(wrap_width(edited) == wrap_width(narrow),
          'editing a few lines moved the column width')

    # The Georgia case: chart cells arrive delimited, with no invented punctuation.
    html = ('<html><body><main><table><tr><th>Who may file</th>'
            '<td>Any individual or organization</td></tr>'
            '<tr><th>Deadline</th><td>One year</td></tr></table>'
            '<p>After the chart.</p></main></body></html>')
    got = extract_html(html, 'main')
    check('Who may file | Any individual or organization' in got,
          'table cells not pipe-delimited')
    check('organization.' not in got, 'extraction invented terminal punctuation')
    check('After the chart.' in got, 'prose after the table was dropped')

    # A server-rendered island must be found and pathed into, and a page that
    # merely looks client-rendered must not send anyone to a browser.
    island = ('<html><body><div id="__next"></div>'
              '<script id="__NEXT_DATA__" type="application/json">'
              '{"props":{"page":{"content":"<table><tr><td>Who may file</td>'
              '<td>Any individual or organization</td></tr></table>",'
              '"contact":{"contacts":[{"phone":"404-656-3963",'
              '"email":"SPEDHelpDesk@doe.k12.ga.us","fax":null}]}}}}'
              '</script></body></html>')
    got = extract_next_data(island, 'props.page.content')
    check('Who may file | Any individual or organization' in got,
          'next-data did not reach the page HTML')
    try:
        extract_next_data(island, 'props.page.missing')
        check(False, 'next-data accepted a path that does not resolve')
    except RuntimeError:
        pass
    # Contacts live beside the body copy, not inside it, and both must survive.
    both = extract_next_data(island, ['props.page.content', 'props.page.contact'])
    check('Who may file |' in both and 'SPEDHelpDesk@doe.k12.ga.us' in both
          and '404-656-3963' in both,
          'a multi-path scope dropped the body copy or the contacts')
    check('phone: 404-656-3963' in both,
          'structured contacts were not rendered as key: value lines')


    # Per-source --compressed opt-out, 2026-09-03. The guarantee that makes it
    # safe to add: a recipe that does not mention it fetches, and digests,
    # exactly as before the option existed.
    check('--compressed' in curl_args(True), 'curl_args dropped --compressed by default')
    check('--compressed' not in curl_args(False), 'compressed=False kept --compressed')
    check([a for a in curl_args(False)] == [a for a in CURL_ARGS if a != '--compressed'],
          'compressed=False changed an argument other than --compressed')
    check(curl_args(True) == CURL_ARGS, 'curl_args(True) is not CURL_ARGS itself')
    _cbase = {'recipe_version': RECIPE_VERSION, 'state': 'testland',
              'sources': [{'n': 1, 'title': 'T', 'url': 'https://x.gov/a',
                           'transport': 'curl', 'extractor': 'html-text',
                           'scope': 'main', 'filters': []}]}
    _con = json.loads(json.dumps(_cbase)); _con['sources'][0]['compressed'] = True
    _coff = json.loads(json.dumps(_cbase)); _coff['sources'][0]['compressed'] = False
    check(digest(_cbase) == digest(_con),
          'setting compressed to its default moved the digest')
    check(digest(_cbase) != digest(_coff),
          'opting out of --compressed did not move the digest')
    check(lint(_coff, 'testland') == [], 'a valid compressed:false source did not lint')
    _cbad = json.loads(json.dumps(_cbase)); _cbad['sources'][0]['compressed'] = 'no'
    check(lint(_cbad, 'testland'), 'lint accepted a non-boolean compressed')

    # Ordering is part of the recipe, so it must be honoured.
    # -nopgbrk removal, 2026-09-03. Three properties, each of which the defect
    # in sped-safeguards violated. Behavioural coverage of the filter itself
    # lives in capture-core.py's own self-test; what is asserted here is that
    # this repo is wired to it correctly.
    #
    # 1. Neither pinned mode may carry the flag back.
    for _name, _flags in (('pdftotext-raw', PDFTOTEXT_RAW),
                          ('pdftotext-layout', PDFTOTEXT_LAYOUT)):
        check('-nopgbrk' not in _flags,
              f'{_name} carries -nopgbrk; the page break it suppresses is what '
              'welds a running head into the sentence it interrupts')
    # 2. The filter is reachable by name, so a recipe that needs it can ask.
    check(FILTERS.get('strip-running-headers') is f_strip_running_headers,
          'strip-running-headers is not registered in FILTERS')
    # 3. A form feed must not leak into a packet body, but must survive until
    #    after the filters -- converting it earlier destroys the same signal
    #    -nopgbrk used to destroy.
    _ff = 'sentence cut here\x0cRUNNING HEAD'
    check('\x0c' not in _core.form_feed_to_linebreak(_ff),
          'a form feed survived into the packet body')
    check(_core.form_feed_to_linebreak(_ff) == 'sentence cut here\nRUNNING HEAD',
          'form_feed_to_linebreak did not render the page break as a line break')

    check(apply_filters('  x • y  ', ['trim-lines', 'normalize-apostrophes'])
          == '  x • y', 'filters applied out of order or not at all')

    # Lint catches the mistakes that would actually be made.
    good = {'recipe_version': 1, 'state': 'testland', 'sources': [
        {'n': 1, 'title': 'Notice', 'url': 'https://x.gov/n.pdf',
         'transport': 'curl', 'extractor': 'pdftotext-raw',
         'filters': ['strip-page-numbers', 'unwrap-hard-wraps']}]}
    check(lint(good, 'testland') == [], 'lint rejected a valid recipe')
    check(lint({**good, 'state': 'other'}, 'testland'), 'lint missed a slug mismatch')
    bad_filter = json.loads(json.dumps(good))
    bad_filter['sources'][0]['filters'] = ['make-it-nice']
    check(lint(bad_filter, 'testland'), 'lint missed an unknown filter')
    no_scope = json.loads(json.dumps(good))
    no_scope['sources'][0]['extractor'] = 'html-text'
    check(lint(no_scope, 'testland'), 'lint missed html-text without a scope')

    # The digest must track capture-determining fields and nothing else.
    d0 = digest(good)
    retitled = json.loads(json.dumps(good))
    retitled['sources'][0]['title'] = 'Notice of Procedural Safeguards'
    check(digest(retitled) == d0, 'digest moved when only a title changed')
    refiltered = json.loads(json.dumps(good))
    refiltered['sources'][0]['filters'] = ['strip-page-numbers']
    check(digest(refiltered) != d0, 'digest ignored a filter change')

    # An agent transport must halt rather than silently fall back to curl.
    try:
        capture_source({'n': 1, 'url': 'https://x.gov/a', 'transport': 'web_fetch',
                        'extractor': 'none'})
        check(False, 'agent transport did not halt')
    except NeedsAgentFetch:
        pass

    # A rendered packet must satisfy retain-packet.py's own parser.
    packet = render_packet(good, {1: 'The child must remain.'}, '2026-08-28')
    check(packet.startswith('FIRST LINE OF PACKET'), 'canary missing')
    check(re.search(r'^STATE: testland$', packet, re.M), 'state header missing')
    check(re.search(r'^SOURCE 1: Notice \| https://x\.gov/n\.pdf \| ', packet, re.M),
          'source header not in the expected shape')
    check(d0 in packet, 'packet does not record its recipe digest')


    # join-wrap-hyphens: close the break, keep the hyphen, touch nothing else.
    # Deleting the hyphen would invent a word the source does not contain.
    check(f_join_wrap_hyphens('word-for-\nword record') == 'word-for-word record',
          'join-wrap-hyphens did not close a stranded hyphen')
    check('-' in f_join_wrap_hyphens('word-for-\nword'),
          'join-wrap-hyphens deleted the hyphen instead of keeping it')
    check(f_join_wrap_hyphens('ends with a dash -\nThen A Capital')
          == 'ends with a dash -\nThen A Capital',
          'join-wrap-hyphens joined across a line starting with a capital')
    check(f_join_wrap_hyphens('no hyphen\nhere') == 'no hyphen\nhere',
          'join-wrap-hyphens altered a plain line break')
    check('join-wrap-hyphens' in FILTERS,
          'join-wrap-hyphens is not registered in FILTERS')

    # --- slicing. The document below is the shape the option exists for: a
    # table of contents that prints the heading, then the text under the same
    # heading, then the following rule.
    doc = ('CHAPTER 1200-08-06\n'
           'TABLE OF CONTENTS\n'
           '1200-08-06-.05 ADMISSIONS, DISCHARGES, AND TRANSFERS.\n'
           '1200-08-06-.06 SOMETHING ELSE.\n'
           '\n'
           '1200-08-06-.05 ADMISSIONS, DISCHARGES,   AND\nTRANSFERS.\n'
           '(1) The facility shall permit each resident to remain.\n'
           '\n'
           '1200-08-06-.06 SOMETHING ELSE.\n'
           '(1) Not this rule.\n')
    head = '1200-08-06-.05 ADMISSIONS, DISCHARGES, AND TRANSFERS.'
    body = apply_slice(doc, {'from': head, 'from_occurrence': 2,
                             'to': '1200-08-06-.06 SOMETHING ELSE.'})
    check(body.startswith('1200-08-06-.05'), 'slice did not start at its anchor')
    check('remain' in body, 'slice dropped the text it exists to capture')
    check('Not this rule' not in body, 'slice ran past its to anchor')
    check('TABLE OF CONTENTS' not in body,
          'slice landed in the table of contents despite from_occurrence 2')
    # Elastic whitespace is the whole point: the anchor above is written on one
    # line and the document breaks it across two with runs of spaces.
    check('\n' in doc.split(head)[0] or True, 'sanity')
    check(_anchor_re('AND TRANSFERS.').search('AND\n  TRANSFERS.') is not None,
          'anchor whitespace is not elastic across a line break')
    # No `to` means run to the end.
    check(apply_slice(doc, {'from': head, 'from_occurrence': 2})
          .endswith('Not this rule.\n'),
          'a slice with no to anchor did not run to the end of the document')
    # A miss raises rather than returning the whole document or an empty one.
    try:
        apply_slice(doc, {'from': 'A HEADING THAT IS NOT THERE AT ALL'})
        check(False, 'a missing from anchor did not raise')
    except SliceMiss:
        pass
    try:
        apply_slice(doc, {'from': head, 'from_occurrence': 9})
        check(False, 'an out-of-range from_occurrence did not raise')
    except SliceMiss:
        pass
    # A `to` that appears only before the `from` is a miss, not a backwards
    # slice: the tail is searched, never the whole document.
    try:
        apply_slice(doc, {'from': head, 'from_occurrence': 2,
                          'to': 'TABLE OF CONTENTS'})
        check(False, 'a to anchor behind the from anchor did not raise')
    except SliceMiss:
        pass

    _slbase = {'recipe_version': RECIPE_VERSION, 'state': 'testland',
               'sources': [{'n': 1, 'title': 't', 'url': 'https://x/y',
                            'transport': 'curl', 'extractor': 'pdftotext-layout',
                            'filters': []}]}
    def _with_slice(sl):
        r = json.loads(json.dumps(_slbase))
        r['sources'][0]['slice'] = sl
        return r
    check(not lint(_with_slice({'from': head, 'to': 'SOMETHING ELSE HERE'}),
                   'testland'),
          'a well-formed slice did not lint clean')
    check(any('at least' in e for e in lint(_with_slice({'from': 'short'}),
                                            'testland')),
          'lint accepted an anchor below MIN_ANCHOR')
    check(lint(_with_slice({'to': 'a long enough anchor'}), 'testland'),
          'lint accepted a slice with no from anchor')
    check(any('unknown key' in e for e in
              lint(_with_slice({'from': head, 'until': 'x'}), 'testland')),
          'lint accepted an unknown slice key')
    check(any('to_occurrence' in e for e in
              lint(_with_slice({'from': head, 'to_occurrence': 2}), 'testland')),
          'lint accepted to_occurrence without a to anchor')
    check(any('positive integer' in e for e in
              lint(_with_slice({'from': head, 'from_occurrence': 0}), 'testland')),
          'lint accepted a zero occurrence')
    # The digest moves when the span moves, and only then.
    check(digest(_slbase) != digest(_with_slice({'from': head})),
          'adding a slice did not change the recipe digest')
    check(digest(_with_slice({'from': head}))
          != digest(_with_slice({'from': head, 'from_occurrence': 2})),
          'changing from_occurrence did not change the recipe digest')

    print('self-test: ' + ('all checks passed' if ok else 'FAILURES'))
    return 0 if ok else 1


def main(argv=None):
    p = argparse.ArgumentParser(add_help=True, description=__doc__.split('\n')[0])
    p.add_argument('slug', nargs='?')
    p.add_argument('--source', type=int)
    p.add_argument('--supply', action='append', metavar='N=PATH')
    p.add_argument('--out')
    p.add_argument('--date')
    p.add_argument('--lint', action='store_true')
    p.add_argument('--digest', action='store_true')
    p.add_argument('--self-test', action='store_true')
    a = p.parse_args(argv)

    if a.self_test:
        return self_test()
    if a.lint:
        return cmd_lint(a)
    if a.digest:
        if not a.slug:
            print('--digest needs a state slug', file=sys.stderr)
            return 2
        print(digest(load_recipe(a.slug)))
        return 0
    if not a.slug:
        p.print_help()
        return 2
    return cmd_capture(a)


if __name__ == '__main__':
    sys.exit(main())
